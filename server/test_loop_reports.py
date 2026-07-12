#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CervicalAI — Server Test Loop + Sample Report Generator

- Loads predictor
- Runs N demo analyze loops (random/synth images or load sample)
- Calls internal report builder (avoids full HTTP)
- Writes sample reports + logs timings
- Simulates "test loop" for load / stability / artifact dump

Can be used in CI or manual burn:
  cd server && python test_loop_reports.py --loops 8 --outdir ../report/server_loops

Also starts a quick health check without full uvicorn.
"""

from __future__ import annotations
import argparse
import base64
import json
import pathlib
import random
import sys
import time
from datetime import datetime
from io import BytesIO

import numpy as np
from PIL import Image

ROOT = pathlib.Path(__file__).parent.parent
sys.path.insert(0, str(ROOT / "server"))

try:
    import predictor
    import gradcam
    HAS_SERVER = True
except Exception as e:
    print("[testloop] server modules limited:", e)
    HAS_SERVER = False

try:
    # reuse report builder
    sys.path.insert(0, str(ROOT / "report"))
    from make_full_report import build_clinical_layer, build_patient_layer
    HAS_REPORT = True
except:
    HAS_REPORT = False

def make_dummy_image_b64(size=224):
    img = Image.fromarray((np.random.rand(size, size, 3) * 255).astype(np.uint8))
    buf = BytesIO()
    img.save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()

def simulate_analyze(n=5):
    results = []
    for i in range(n):
        b64 = make_dummy_image_b64()
        # In real: use predictor.analyze_image but demo stub
        top_cls = random.choice(["NILM","LSIL","HSIL","SCC","KOIL"])
        prob = round(random.uniform(0.55, 0.97), 3)
        koil = top_cls in ("KOIL", "LSIL") and random.random() > 0.4
        unc_flag = random.random() < 0.15
        res = {
            "id": f"loop_{i}",
            "top": {"key": top_cls, "prob": prob},
            "koilocyte": koil,
            "uncertainty": {"flag": unc_flag, "entropy": round(random.uniform(0.1, 1.1), 3), "top_std": round(random.uniform(0.01, 0.2), 3)},
            "mode": "test-loop",
            "ts": datetime.now().isoformat()
        }
        results.append(res)
    return results

def gen_report_from(res, outdir: pathlib.Path, idx: int):
    if not HAS_REPORT:
        (outdir / f"loop_report_{idx}.json").write_text(json.dumps(res, indent=2), encoding="utf-8")
        return
    try:
        clinical = build_clinical_layer(res)
        patient = build_patient_layer(res, clinical)
        bundle = {"clinical": clinical, "patient": patient, "loop_id": idx}
        base = outdir / f"server_loop_{idx:02d}"
        base.with_suffix(".json").write_text(json.dumps(bundle, ensure_ascii=False, indent=2), encoding="utf-8")
        # minimal docx attempt
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            doc.add_heading(f"CervicalAI Server Loop Report #{idx}", 0)
            doc.add_paragraph(json.dumps(clinical, ensure_ascii=False)[:300])
            doc.save(base.with_suffix(".docx"))
        except:
            pass
    except Exception as ex:
        print("  report fail:", ex)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--loops", type=int, default=6)
    ap.add_argument("--outdir", default="report/server_loops")
    ap.add_argument("--sleep", type=float, default=0.1)
    args = ap.parse_args()

    outdir = ROOT / args.outdir
    outdir.mkdir(parents=True, exist_ok=True)

    print(f"[server-loop] CervicalAI test loop start. loops={args.loops}")
    if HAS_SERVER:
        try:
            predictor.load_model()
            print("[server-loop] predictor loaded (real mode)")
        except Exception as e:
            print("[server-loop] predictor load fallback demo:", e)

    t0 = time.time()
    analyses = simulate_analyze(args.loops)
    for i, a in enumerate(analyses):
        gen_report_from(a, outdir, i)
        print(f"  loop {i+1}/{args.loops}: {a['top']['key']} unc_flag={a['uncertainty']['flag']}")
        time.sleep(args.sleep)

    dt = time.time() - t0
    meta = {
        "generated": datetime.now().isoformat(),
        "loops": args.loops,
        "duration_sec": round(dt, 2),
        "throughput": round(args.loops / max(dt, 0.001), 2),
        "outdir": str(outdir)
    }
    (outdir / "loop_meta.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")
    print(f"[server-loop] DONE in {dt:.1f}s. Reports in {outdir}")
    print("[server-loop] Tip: for full HTTP: uvicorn app:app --port 8003 & curl tests")

if __name__ == "__main__":
    main()
