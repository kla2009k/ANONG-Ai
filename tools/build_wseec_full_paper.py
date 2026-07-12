from __future__ import annotations

import json
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from build_full_booklet import ASSETS, make_charts


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "wseec_2026"
DOCX_OUT = OUT_DIR / "CerviCo_Pilot_WSEEC_2026_Full_Paper.docx"
PDF_OUT = OUT_DIR / "CerviCo_Pilot_WSEEC_2026_Full_Paper.pdf"
AUDIT_OUT = OUT_DIR / "WSEEC_2026_FULL_PAPER_AUDIT.json"

TITLE = (
    "DEVELOPMENT OF AN EXPLAINABLE AND UNCERTAINTY-AWARE AI SYSTEM FOR "
    "CERVICAL CELL ABNORMALITY SCREENING AND HPV-RELATED MORPHOLOGY RISK "
    "ASSESSMENT FROM THINPREP-STYLE IMAGES"
)

AUTHORS = [
    "AUTHOR 1: ______________________________",
    "AUTHOR 2: ______________________________",
    "AUTHOR 3: ______________________________",
    "AUTHOR 4: ______________________________",
    "AUTHOR 5: ______________________________",
    "AUTHOR 6: ______________________________",
]

ABSTRACT = (
    "Cervical cancer is largely preventable when screening, interpretation, communication, "
    "and follow-up operate as a continuous pathway. This study developed CerviCo-Pilot, an "
    "explainable and uncertainty-aware research prototype for clinician-in-the-loop cervical "
    "cytology screening support. The system combines a Bethesda-style five-class output "
    "(NILM, LSIL, HSIL, SCC, and a KOIL placeholder), a binary normal/abnormal safety-triage "
    "view, Grad-CAM visual review support, Monte Carlo Dropout uncertainty, and a report lock "
    "that requires clinician sign-off before patient-facing communication. An EfficientNet-B0 "
    "model was trained and evaluated using 917 real images from the public Herlev cervical "
    "cytology dataset; segmentation masks were excluded from performance evidence. On the "
    "held-out test set (n=137), five-class accuracy was 0.6934, macro F1 was 0.5545, macro "
    "AUROC was 0.7311, and quadratic weighted kappa was 0.687. The binary safety-triage view "
    "achieved sensitivity 1.000, specificity 0.7222, AUROC 0.964, AUPRC 0.9856, and a "
    "confusion matrix of TP=101, TN=26, FP=10, and FN=0. Five-fold cross-validation produced "
    "binary sensitivity of 0.9867 +/- 0.0086 and AUROC of 0.9435 +/- 0.0448. Post-hoc "
    "temperature scaling improved held-out calibration metrics without materially changing "
    "discrimination. The prototype demonstrates a complete review workflow rather than an "
    "autonomous diagnostic device. Its present evidence is limited to the Herlev domain; it "
    "does not replace molecular HPV testing, and Thai ThinPrep validation, paired HPV "
    "endpoints, reader studies, and prospective evaluation remain necessary before clinical "
    "use."
)

KEYWORDS = [
    "cervical cytology",
    "explainable artificial intelligence",
    "uncertainty-aware triage",
    "Grad-CAM",
    "clinician-in-the-loop",
    "ThinPrep",
]

REFERENCES = [
    "[1] World Health Organization, Cervical cancer, WHO Fact Sheet, 2026. Available: https://www.who.int/news-room/fact-sheets/detail/cervical-cancer.",
    "[2] World Health Organization, Global Strategy to Accelerate the Elimination of Cervical Cancer as a Public Health Problem, WHO, 2020.",
    "[3] Centers for Disease Control and Prevention, Screening for Cervical Cancer, CDC, 2025. Available: https://www.cdc.gov/cervical-cancer/screening/.",
    "[4] R. Nayar and D. C. Wilbur, Eds., The Bethesda System for Reporting Cervical Cytology, 3rd ed., Springer, 2015.",
    "[5] J. Jantzen, J. Norup, G. Dounias, and B. Bjerregaard, Pap-smear benchmark data for pattern classification, NiSIS 2005, 2005.",
    "[6] M. Tan and Q. V. Le, EfficientNet: Rethinking model scaling for convolutional neural networks, Proc. ICML, 2019, pp. 6105-6114.",
    "[7] R. R. Selvaraju et al., Grad-CAM: Visual explanations from deep networks via gradient-based localization, Proc. ICCV, 2017, pp. 618-626.",
    "[8] Y. Gal and Z. Ghahramani, Dropout as a Bayesian approximation: Representing model uncertainty in deep learning, Proc. ICML, 2016, pp. 1050-1059.",
    "[9] C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, On calibration of modern neural networks, Proc. ICML, 2017, pp. 1321-1330.",
    "[10] U.S. FDA, Health Canada, and MHRA, Good Machine Learning Practice for Medical Device Development: Guiding Principles, 2021.",
    "[11] X. Liu et al., Reporting guidelines for clinical trial reports for interventions involving artificial intelligence: the CONSORT-AI extension, Nature Medicine, vol. 26, pp. 1364-1374, 2020.",
    "[12] B. Vasey et al., Reporting guideline for the early-stage clinical evaluation of decision support systems driven by artificial intelligence: DECIDE-AI, Nature Medicine, vol. 28, pp. 924-933, 2022.",
    "[13] S. G. Finlayson et al., The clinician and dataset shift in artificial intelligence, New England Journal of Medicine, vol. 385, pp. 283-286, 2021.",
]


def P(text: str) -> dict:
    return {"type": "p", "text": text}


def H(text: str) -> dict:
    return {"type": "heading", "text": text}


def B(items: list[str]) -> dict:
    return {"type": "bullets", "items": items}


def T(caption: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> dict:
    return {"type": "table", "caption": caption, "headers": headers, "rows": rows, "widths": widths}


def F(path: Path, caption: str, width_cm: float, height_cm: float | None = None) -> dict:
    return {"type": "figure", "path": path, "caption": caption, "width_cm": width_cm, "height_cm": height_cm}


def N(text: str) -> dict:
    return {"type": "note", "text": text}


PAGES: list[dict] = [
    {"kind": "cover"},
    {"kind": "toc"},
    {"kind": "abstract"},
    {
        "kind": "chapter", "chapter": 1, "title": "INTRODUCTION", "elements": [
            H("1.1 Research Background"),
            P("Cervical cancer remains an important public-health problem, yet it can be prevented or treated early when vaccination, screening, diagnosis, and follow-up are connected [1,2]. Cytology-based screening, including Pap and liquid-based preparations such as ThinPrep, examines cellular morphology, while HPV molecular testing detects viral material; the two modalities must not be described as interchangeable [3,4]."),
            P("The practical gap is not limited to image classification. Screening programs also depend on timely review, understandable communication, and reliable referral. In settings with limited cytology expertise, a compact offline decision-support tool may help prioritize abnormal images, expose the image regions that influenced a prediction, and return uncertain cases to human review."),
            H("1.2 Problem Statement"),
            P("Many image classifiers provide only a label and confidence score. Such output is insufficient for a safety-sensitive workflow because confidence may be miscalibrated, explanations may be over-trusted, and a patient-facing report may be released without appropriate review. A useful prototype should therefore combine grading, triage, explainability, uncertainty, and explicit clinician control."),
            H("1.3 Objectives and Significance"),
            B([
                "Develop a Bethesda-style five-class cervical cytology screening-support model.",
                "Add a high-sensitivity binary safety-triage view without replacing the five-class output.",
                "Provide Grad-CAM review support and uncertainty-aware abstention.",
                "Demonstrate clinician sign-off, patient-report locking, and offline operation.",
                "Define an evidence boundary for future Thai ThinPrep and paired HPV validation.",
            ]),
            N("Current scope: retrospective public-dataset research prototype; not an autonomous diagnostic device and not a molecular HPV test."),
        ],
    },
    {
        "kind": "chapter", "chapter": 2, "title": "LITERATURE REVIEW", "elements": [
            H("2.1 Cervical Cytology and the Bethesda Framework"),
            P("The Bethesda System standardizes cervical cytology terminology and separates negative findings from low-grade, high-grade, and malignant squamous abnormalities [4]. Preserving multiple output classes is useful for cytology workflow communication, whereas a secondary binary view is useful for screening-safety analysis."),
            H("2.2 Image Classification, Explainability, and Uncertainty"),
            P("The Herlev benchmark contains 917 single-cell cervical cytology images and has been widely used for pattern-classification research [5]. EfficientNet scales network depth, width, and resolution in a computationally efficient manner [6], making EfficientNet-B0 suitable for an offline prototype. Grad-CAM localizes image regions associated with a class score [7], but it should be treated as a review aid rather than causal proof. Monte Carlo Dropout offers an approximate measure of predictive uncertainty [8]."),
            H("2.3 Calibration and Medical-AI Governance"),
            P("Modern neural networks may be overconfident; temperature scaling is a simple post-hoc calibration method that preserves class ordering [9]. Medical-AI guidance emphasizes representative data, human-AI interaction, lifecycle monitoring, and transparent reporting [10-12]. Dataset shift remains a central risk when moving from a public benchmark to a different population, preparation protocol, microscope, or camera [13]."),
            T(
                "Table 1. Design principles derived from the literature.",
                ["Principle", "Implementation in CerviCo-Pilot", "Boundary"],
                [
                    ["Structured cytology", "Five-class Bethesda-style suggestion", "Not a final cytology report"],
                    ["Screening safety", "Binary normal/abnormal triage view", "Not the sole product output"],
                    ["Explainability", "Grad-CAM heatmap", "Review aid, not causal proof"],
                    ["Uncertainty", "MC Dropout flag and abstention", "Requires external validation"],
                    ["Governance", "Clinician sign-off and report lock", "Demo workflow, not regulated software"],
                ],
                [3.2, 6.0, 4.8],
            ),
        ],
    },
    {
        "kind": "chapter", "chapter": 3, "title": "RESEARCH METHOD", "elements": [
            H("3.1 Study Design and Dataset"),
            P("This was a retrospective model-development and prototype-integration study. The canonical evidence used 917 real images from the public Herlev cervical cytology dataset [5]. Segmentation-mask files and synthetic-heavy historical experiments were excluded from public performance evidence. The held-out test set contained 137 images."),
            T(
                "Table 2. Dataset and model card summary.",
                ["Item", "Specification"],
                [
                    ["Dataset", "Herlev public cervical cytology dataset"],
                    ["Images", "917 real images; masks excluded"],
                    ["Architecture", "EfficientNet-B0 with ImageNet transfer learning"],
                    ["Five-class output", "NILM, LSIL, HSIL, SCC, KOIL placeholder"],
                    ["Binary triage", "NILM=normal; LSIL/HSIL/SCC/KOIL=abnormal"],
                    ["Evidence status", "Phase 1 Herlev-only research prototype"],
                ],
                [4.0, 10.0],
            ),
            H("3.2 End-to-End Workflow"),
            P("The workflow accepts a cytology image, produces five-class probabilities and a binary triage interpretation, generates a Grad-CAM overlay, estimates uncertainty, and routes the result to clinician confirmation, correction, or rejection. A patient-facing report remains locked when the case is uncertain or has not been signed off."),
            F(ASSETS / "workflow_diagram.png", "Figure 1. CerviCo-Pilot clinician-in-the-loop workflow.", 13.2, 4.7),
        ],
    },
    {
        "kind": "chapter_cont", "chapter": 3, "title": "RESEARCH METHOD (CONTINUED)", "elements": [
            H("3.3 Model Development"),
            P("Images were resized for EfficientNet-B0 input and processed using the project training pipeline. Transfer learning, class-aware training, focal-loss experiments, oversampling, and test-time augmentation were used to address limited data and class imbalance. The shipped checkpoint is models/best_cervical.pt; only metrics reproduced from this checkpoint are used in this paper."),
            H("3.4 Evaluation Strategy"),
            P("Five-class evaluation included accuracy, macro F1, macro AUROC, per-class recall, and quadratic weighted kappa. The binary safety-triage view included sensitivity, specificity, AUROC, AUPRC, Matthews correlation coefficient, and confusion counts. Uncertainty was assessed with repeated stochastic forward passes. Evaluation used a held-out test set, bootstrap confidence intervals, and five-fold cross-validation."),
            T(
                "Table 3. Primary evaluation measures.",
                ["Layer", "Measures", "Purpose"],
                [
                    ["Five-class", "Accuracy, macro F1, macro AUROC, QWK", "Detailed grading performance"],
                    ["Binary triage", "Sensitivity, specificity, AUROC, AUPRC, MCC", "Screening-safety view"],
                    ["Calibration", "ECE and Brier score", "Probability reliability"],
                    ["Error analysis", "Confusion matrix and case gallery", "Identify class-specific failure modes"],
                    ["Cross-validation", "Five folds, mean +/- SD", "Assess split sensitivity"],
                ],
                [3.0, 6.4, 4.6],
            ),
            H("3.5 Post-Hoc Calibration"),
            P("A scalar temperature was fitted after training. Calibration was compared before and after scaling on the held-out Herlev test set. This experiment was designed to improve probability communication within the current domain; it does not establish calibration on Thai ThinPrep images."),
        ],
    },
    {
        "kind": "chapter_cont", "chapter": 3, "title": "RESEARCH METHOD (CONTINUED)", "elements": [
            H("3.6 Explainability and Uncertainty Policy"),
            P("Grad-CAM was generated from the classifier to show the region associated with the predicted class [7]. The interface displays the original image, heatmap, class probabilities, and uncertainty. High uncertainty activates abstention language and blocks patient-report release. The heatmap is never presented as proof that the model has identified a biological cause."),
            F(ASSETS / "sample_gradcam_grid.jpg", "Figure 2. Herlev cytology examples with Grad-CAM review overlays.", 13.0, 6.1),
            H("3.7 Prototype System and Safety Controls"),
            P("The prototype combines a React interface and FastAPI service. It supports sample review, image upload when the local server is active, clinician confirm/edit/reject actions, a report preview, model card, evidence page, and audit demonstration. The audit is explicitly labeled as a demo and does not satisfy regulated clinical traceability requirements."),
            B([
                "Intended users: supervised evaluators and trained clinical personnel in research settings.",
                "Do not use: autonomous diagnosis, molecular HPV status, or unsupervised patient release.",
                "Deployment boundary: real authentication, de-identification, governance, and external validation are not yet complete.",
            ]),
        ],
    },
    {
        "kind": "chapter", "chapter": 4, "title": "RESULTS AND DISCUSSION", "elements": [
            H("4.1 Primary Quantitative Results"),
            T(
                "Table 4. Held-out and cross-validation results.",
                ["Measure", "Result", "Interpretation"],
                [
                    ["Held-out five-class accuracy", "0.6934", "Moderate detailed grading"],
                    ["Held-out macro F1", "0.5545", "Class-specific limitations remain"],
                    ["Held-out macro AUROC", "0.7311", "Moderate multiclass discrimination"],
                    ["Held-out QWK", "0.687", "Substantial severity agreement"],
                    ["Held-out binary sensitivity", "1.000", "FN=0 on this held-out split"],
                    ["Held-out binary specificity", "0.7222", "Some over-referral"],
                    ["Held-out binary AUROC / AUPRC", "0.964 / 0.9856", "Strong Herlev-domain triage"],
                    ["Five-fold binary sensitivity", "0.9867 +/- 0.0086", "High sensitivity across folds"],
                    ["Five-fold binary AUROC", "0.9435 +/- 0.0448", "Robust within-domain discrimination"],
                ],
                [5.0, 3.4, 5.6],
            ),
            F(ASSETS / "headline_metrics.png", "Figure 3. Headline held-out performance metrics.", 12.6, 6.0),
            P("The central finding is the difference between the two output layers. The five-class task retained clinically useful detail but remained moderate, whereas the binary safety-triage view achieved high sensitivity. Therefore, the binary result should be interpreted as a safety layer rather than a replacement for the five-class objective."),
        ],
    },
    {
        "kind": "chapter_cont", "chapter": 4, "title": "RESULTS AND DISCUSSION (CONTINUED)", "elements": [
            H("4.2 Class-Specific Error Analysis"),
            F(ASSETS / "confusion_matrix.png", "Figure 4. Five-class output-space confusion matrix on the held-out test set; KOIL recall is not estimable because support is zero.", 9.2, 7.4),
            T(
                "Table 5. Held-out per-class recall and evidence limitations.",
                ["Class", "Support", "Recall", "Interpretation"],
                [
                    ["NILM", "36", "0.7222", "False positives reduce specificity"],
                    ["LSIL", "49", "0.6122", "Overlap with higher-grade morphology"],
                    ["HSIL", "30", "0.8667", "Strongest high-grade recall"],
                    ["SCC", "22", "0.5909", "Requires improvement"],
                    ["KOIL", "0", "N/A", "Not estimable; no true KOIL support in Herlev"],
                ],
                [2.5, 2.4, 2.4, 6.7],
            ),
            P("The absence of true KOIL examples is a structural evidence gap, not a successful negative finding. KOIL is retained only as a future product target and interface placeholder. SCC recall also remains insufficient for any claim of autonomous grading."),
        ],
    },
    {
        "kind": "chapter_cont", "chapter": 4, "title": "RESULTS AND DISCUSSION (CONTINUED)", "elements": [
            H("4.3 Calibration and Uncertainty"),
            T(
                "Table 6. Held-out calibration before and after temperature scaling.",
                ["Measure", "Before", "After"],
                [
                    ["Multiclass ECE", "0.066853", "0.038670"],
                    ["Binary ECE", "0.090679", "0.073787"],
                    ["Binary Brier score", "0.071015", "0.066994"],
                    ["Binary AUROC", "0.963971", "0.963421"],
                ],
                [6.0, 4.0, 4.0],
            ),
            F(ASSETS / "calibration_before_after.png", "Figure 5. Calibration comparison after post-hoc temperature scaling.", 12.5, 5.8),
            H("4.4 System-Level Discussion"),
            P("Temperature scaling improved held-out ECE and Brier score while leaving discrimination nearly unchanged, which is consistent with the purpose of post-hoc calibration [9]. Nevertheless, calibration is domain-dependent. The result cannot be extended to Thai ThinPrep images without external data."),
            P("The prototype's contribution is the integration of quantitative evaluation with a review workflow: Grad-CAM can be challenged, uncertainty can trigger abstention, and the patient report is gated by clinician action. This responds to human-AI interaction and transparent-reporting principles [10-12]. The principal remaining risk is dataset shift [13]."),
            N("No clinical-utility claim is made. A reader study and prospective workflow evaluation are required to test whether the system improves sensitivity, review time, or follow-up without causing automation bias."),
        ],
    },
    {
        "kind": "chapter", "chapter": 5, "title": "CONCLUSION", "elements": [
            H("5.1 Conclusion"),
            P("CerviCo-Pilot demonstrates an end-to-end, explainable, and uncertainty-aware cervical cytology screening-support prototype. On real Herlev images, the binary safety-triage layer achieved held-out sensitivity of 1.000 and AUROC of 0.964, while the five-class task achieved accuracy of 0.6934 and macro AUROC of 0.7311. These results support continued research but do not justify autonomous diagnosis or deployment."),
            P("The main contribution is the system design: detailed cytology output is preserved, a binary view emphasizes screening safety, Grad-CAM supports visual review, uncertainty enables abstention, and clinician sign-off controls patient-facing communication. The system can operate locally and present its limitations alongside its predictions."),
            H("5.2 Limitations"),
            B([
                "Evidence is restricted to a small public single-cell dataset.",
                "No Thai ThinPrep/LBC external validation has been completed.",
                "No paired HPV DNA/RNA endpoint is available; HPV output is morphology-related only.",
                "KOIL has no true support in the current dataset, and SCC recall remains limited.",
                "Grad-CAM, uncertainty, and report wording have not yet been evaluated in a clinical reader study.",
            ]),
            H("5.3 Future Work"),
            P("The next phase will collect de-identified Thai ThinPrep/LBC images under ethics and data-governance approval, lock an external test set, add paired molecular HPV labels, perform pathologist-reviewed error analysis, and conduct an unaided-versus-AI-assisted reader study. A prospective pilot should measure workflow outcomes such as review time, notification time, follow-up completion, and over-reliance."),
            H("5.4 Research Integrity Statement"),
            P("All performance values in this paper were copied from canonical project evaluation files for the shipped EfficientNet-B0 checkpoint. External sources support background and methodology only; they do not create performance evidence for CerviCo-Pilot."),
        ],
    },
]


def set_docx_font(run, size: float = 12, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_docx_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_docx_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_docx_text(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, size=12, after=0, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_docx_font(r, size=size, bold=bold, italic=italic)
    return p


def add_docx_table(doc: Document, element: dict):
    add_docx_text(doc, element["caption"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10.5, after=2)
    table = doc.add_table(rows=1, cols=len(element["headers"]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(element["headers"]):
        cell = table.rows[0].cells[i]
        set_docx_cell_shading(cell, "D9EAF0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_docx_font(r, size=9.5, bold=True)
    for row in element["rows"]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(value)
            set_docx_font(r, size=9.2)
    for row in table.rows:
        for i, width in enumerate(element["widths"]):
            row.cells[i].width = Cm(width)
            set_docx_cell_margins(row.cells[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_docx_figure(doc: Document, element: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run()
    kwargs = {"width": Cm(element["width_cm"])}
    if element.get("height_cm"):
        kwargs["height"] = Cm(element["height_cm"])
    r.add_picture(str(element["path"]), **kwargs)
    add_docx_text(doc, element["caption"], align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10.5, after=2)


def render_docx_element(doc: Document, element: dict):
    kind = element["type"]
    if kind == "p":
        add_docx_text(doc, element["text"], after=3)
    elif kind == "heading":
        add_docx_text(doc, element["text"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=2, before=3)
    elif kind == "bullets":
        for item in element["items"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(item)
            set_docx_font(r, size=12)
    elif kind == "table":
        add_docx_table(doc, element)
    elif kind == "figure":
        add_docx_figure(doc, element)
    elif kind == "note":
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_docx_cell_shading(cell, "F2F2F2")
        set_docx_cell_margins(cell, 80, 100, 80, 100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(element["text"])
        set_docx_font(r, size=10.5, italic=True)


def build_docx() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    for page_index, page in enumerate(PAGES):
        kind = page["kind"]
        if kind == "cover":
            add_docx_text(doc, "FULL PAPER", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=22)
            add_docx_text(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=22)
            logo = doc.add_table(rows=1, cols=1)
            logo.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = logo.cell(0, 0)
            cell.width = Cm(5.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(22)
            p.paragraph_format.space_after = Pt(22)
            r = p.add_run("INSERT INSTITUTION LOGO")
            set_docx_font(r, size=12, bold=True)
            add_docx_text(doc, "COMPILED BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, before=14, after=8)
            for author in AUTHORS:
                add_docx_text(doc, author, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=1)
            add_docx_text(doc, "INSTITUTION NAME: ______________________________", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, before=18, after=3)
            add_docx_text(doc, "CATEGORY: TECHNOLOGY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=3)
            add_docx_text(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=0)
        elif kind == "toc":
            add_docx_text(doc, "TABLE OF CONTENTS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=16)
            toc = [
                ("1. Abstract", "3"),
                ("2. Chapter 1 - Introduction", "4"),
                ("3. Chapter 2 - Literature Review", "5"),
                ("4. Chapter 3 - Research Method", "6"),
                ("5. Chapter 4 - Results and Discussion", "9"),
                ("6. Chapter 5 - Conclusion", "12"),
                ("7. References", "13"),
            ]
            for label, number in toc:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Cm(12.5)
                table.columns[1].width = Cm(1.5)
                p1 = table.cell(0, 0).paragraphs[0]
                p2 = table.cell(0, 1).paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_docx_font(p1.add_run(label), size=12, bold=True)
                set_docx_font(p2.add_run(number), size=12, bold=True)
        elif kind == "abstract":
            add_docx_text(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=12)
            add_docx_text(doc, ABSTRACT, after=10)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.0
            r1 = p.add_run("Keywords: ")
            set_docx_font(r1, size=12, bold=True, italic=True)
            r2 = p.add_run(", ".join(KEYWORDS))
            set_docx_font(r2, size=12, italic=True)
        else:
            add_docx_text(doc, f"CHAPTER {page['chapter']}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=6)
            add_docx_text(doc, page["title"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=12, after=6)
            for element in page["elements"]:
                render_docx_element(doc, element)
        doc.add_page_break()

    add_docx_text(doc, "REFERENCES", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=12, after=10)
    for ref in REFERENCES:
        p = add_docx_text(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, after=5)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)

    core = doc.core_properties
    core.title = TITLE
    core.subject = "WSEEC 2026 Full Paper"
    core.author = "CerviCo-Pilot Team"
    core.keywords = ", ".join(KEYWORDS)
    doc.save(DOCX_OUT)
    return DOCX_OUT


def register_pdf_fonts():
    font_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("Arial", str(font_dir / "arial.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Bold", str(font_dir / "arialbd.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Italic", str(font_dir / "ariali.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-BoldItalic", str(font_dir / "arialbi.ttf")))


def pdf_styles():
    register_pdf_fonts()
    styles = getSampleStyleSheet()
    return {
        "body": ParagraphStyle("BodyArial", parent=styles["BodyText"], fontName="Arial", fontSize=12, leading=14.2, alignment=TA_JUSTIFY, spaceAfter=5),
        "heading": ParagraphStyle("HeadingArial", parent=styles["Heading2"], fontName="Arial-Bold", fontSize=12, leading=14.2, alignment=TA_LEFT, spaceBefore=5, spaceAfter=4),
        "chapter": ParagraphStyle("ChapterArial", parent=styles["Heading1"], fontName="Arial-Bold", fontSize=12, leading=14.2, alignment=TA_CENTER, spaceAfter=6),
        "title": ParagraphStyle("TitleArial", parent=styles["Title"], fontName="Arial-Bold", fontSize=14, leading=17, alignment=TA_CENTER, spaceAfter=12),
        "caption": ParagraphStyle("CaptionArial", parent=styles["BodyText"], fontName="Arial-Italic", fontSize=9.5, leading=11, alignment=TA_CENTER, spaceAfter=4),
        "small": ParagraphStyle("SmallArial", parent=styles["BodyText"], fontName="Arial", fontSize=9.5, leading=11, alignment=TA_LEFT),
        "note": ParagraphStyle("NoteArial", parent=styles["BodyText"], fontName="Arial-Italic", fontSize=10, leading=12, alignment=TA_JUSTIFY),
        "ref": ParagraphStyle("RefArial", parent=styles["BodyText"], fontName="Arial", fontSize=12, leading=14.2, alignment=TA_LEFT, leftIndent=14, firstLineIndent=-14, spaceAfter=5),
    }


def pdf_para(text: str, style):
    return Paragraph(escape(text), style)


def render_pdf_element(story: list, element: dict, styles: dict):
    kind = element["type"]
    if kind == "p":
        story.append(pdf_para(element["text"], styles["body"]))
    elif kind == "heading":
        story.append(pdf_para(element["text"], styles["heading"]))
    elif kind == "bullets":
        for item in element["items"]:
            story.append(pdf_para("- " + item, styles["body"]))
    elif kind == "table":
        data = [[Paragraph(escape(x), styles["small"]) for x in element["headers"]]]
        data.extend([[Paragraph(escape(x), styles["small"]) for x in row] for row in element["rows"]])
        table = Table(data, colWidths=[w * cm for w in element["widths"]], repeatRows=1, hAlign="CENTER")
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF0")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#707070")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(KeepTogether([pdf_para(element["caption"], styles["caption"]), table, Spacer(1, 4)]))
    elif kind == "figure":
        img = Image(str(element["path"]), width=element["width_cm"] * cm, height=(element.get("height_cm") or 6) * cm)
        img.hAlign = "CENTER"
        story.append(KeepTogether([img, pdf_para(element["caption"], styles["caption"])]))
    elif kind == "note":
        note = Table([[pdf_para(element["text"], styles["note"])]], colWidths=[14 * cm])
        note.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F2F2F2")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#888888")),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(note)


def build_pdf() -> Path:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT), pagesize=A4,
        leftMargin=4 * cm, rightMargin=3 * cm,
        topMargin=3 * cm, bottomMargin=3 * cm,
        title=TITLE, author="CerviCo-Pilot Team", subject="WSEEC 2026 Full Paper",
    )
    story = []
    for page in PAGES:
        kind = page["kind"]
        if kind == "cover":
            story.extend([
                pdf_para("FULL PAPER", styles["title"]), Spacer(1, 20),
                pdf_para(TITLE, styles["title"]), Spacer(1, 22),
            ])
            logo = Table([[pdf_para("INSERT INSTITUTION LOGO", styles["chapter"])]], colWidths=[5.2 * cm], rowHeights=[3.3 * cm])
            logo.hAlign = "CENTER"
            logo.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.8, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
            story.extend([logo, Spacer(1, 16), pdf_para("COMPILED BY", styles["title"])])
            for author in AUTHORS:
                story.append(pdf_para(author, styles["chapter"]))
            story.extend([
                Spacer(1, 18),
                pdf_para("INSTITUTION NAME: ______________________________", styles["chapter"]),
                pdf_para("CATEGORY: TECHNOLOGY", styles["chapter"]),
                pdf_para("2026", styles["chapter"]),
            ])
        elif kind == "toc":
            story.append(pdf_para("TABLE OF CONTENTS", styles["title"]))
            toc_data = [
                ["1. Abstract", "3"], ["2. Chapter 1 - Introduction", "4"],
                ["3. Chapter 2 - Literature Review", "5"], ["4. Chapter 3 - Research Method", "6"],
                ["5. Chapter 4 - Results and Discussion", "9"], ["6. Chapter 5 - Conclusion", "12"],
                ["7. References", "13"],
            ]
            table = Table([[pdf_para(a, styles["chapter"]), pdf_para(b, styles["chapter"])] for a, b in toc_data], colWidths=[12.5 * cm, 1.5 * cm])
            table.setStyle(TableStyle([("ALIGN", (1, 0), (1, -1), "RIGHT"), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
            story.append(table)
        elif kind == "abstract":
            story.extend([
                pdf_para("ABSTRACT", styles["title"]),
                pdf_para(ABSTRACT, styles["body"]),
                Spacer(1, 8),
                Paragraph("<b><i>Keywords:</i></b> <i>" + escape(", ".join(KEYWORDS)) + "</i>", styles["body"]),
            ])
        else:
            story.append(pdf_para(f"CHAPTER {page['chapter']}", styles["chapter"]))
            story.append(pdf_para(page["title"], styles["heading"]))
            for element in page["elements"]:
                render_pdf_element(story, element, styles)
        story.append(PageBreak())

    story.append(pdf_para("REFERENCES", styles["heading"]))
    for ref in REFERENCES:
        story.append(pdf_para(ref, styles["ref"]))
    doc.build(story)
    return PDF_OUT


def audit_outputs() -> dict:
    pdf = PdfReader(str(PDF_OUT))
    docx = Document(DOCX_OUT)
    abstract_words = len(ABSTRACT.split())
    audit = {
        "format_source": [
            "docs/wseec_2026/references/Guidebook_WSEEC_2026.pdf",
            "docs/wseec_2026/references/FORMAT_FULL_PAPER.pdf",
        ],
        "requirements": {
            "language": "English",
            "paper_size": "A4",
            "font": "Arial 12 pt",
            "line_spacing": "single",
            "margins_cm": {"left": 4, "right": 3, "top": 3, "bottom": 3},
            "main_page_limit": 12,
            "references_excluded": True,
            "abstract_max_words": 350,
            "keywords_required": "4-6",
        },
        "outputs": {
            "docx": str(DOCX_OUT),
            "pdf": str(PDF_OUT),
            "pdf_pages_total": len(pdf.pages),
            "planned_main_pages": 12,
            "reference_start_page": 13,
            "docx_paragraphs": len(docx.paragraphs),
            "docx_tables": len(docx.tables),
            "docx_figures": len(docx.inline_shapes),
            "abstract_words": abstract_words,
            "keywords": len(KEYWORDS),
        },
        "checks": {
            "abstract_under_350": abstract_words <= 350,
            "keywords_4_to_6": 4 <= len(KEYWORDS) <= 6,
            "pdf_has_references_after_main_pages": len(pdf.pages) >= 13,
            "docx_exists": DOCX_OUT.exists(),
            "pdf_exists": PDF_OUT.exists(),
        },
        "placeholders_to_replace": [
            "AUTHOR 1-6",
            "INSTITUTION LOGO",
            "INSTITUTION NAME",
        ],
    }
    AUDIT_OUT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_charts()
    build_docx()
    build_pdf()
    audit = audit_outputs()
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
