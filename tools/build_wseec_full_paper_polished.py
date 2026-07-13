from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_full_booklet import ASSETS, make_charts
from build_wseec_full_paper import ABSTRACT, AUTHORS, KEYWORDS, REFERENCES, TITLE


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "wseec_2026"
OUT = OUT_DIR / "CerviCo_Pilot_WSEEC_2026_Full_Paper_Polished.docx"
KOIL_ASSETS = ROOT / "models" / "koil_sipakmed" / "evaluation"

TEAL = "0E7490"
NAVY = "143B5E"
INK = "15242E"
MUTED = "526773"
LINE = "B8C8CF"
PALE = "EAF3F6"
PALE_2 = "F5F8F9"
GOLD = "B57B16"
RED = "A74633"


def set_font(run, size=12, bold=None, italic=None, color=INK):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rpr.rFonts.set(qn(f"w:{key}"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE, size="5"):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = node.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            node.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)


def remove_borders(table):
    borders(table, color="FFFFFF", size="0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, color=INK, before=0, after=3):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return para


def chapter_header(doc, number, title, continued=False):
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    cell = band.cell(0, 0)
    cell.width = Cm(14)
    shade(cell, PALE)
    margins(cell, 85, 100, 85, 100)
    borders(band, color=TEAL, size="8")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f"CHAPTER {number}" if str(number).strip() else title
    run = para.add_run(label)
    set_font(run, size=12, bold=True, color=TEAL)
    if str(number).strip():
        p(doc, title + (" (CONTINUED)" if continued else ""), align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, color=NAVY, before=5, after=5)


def heading(doc, text):
    para = p(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, color=NAVY, before=4, after=2)
    para.paragraph_format.keep_with_next = True
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    set_font(run, size=12, color=INK)


def callout(doc, title, text, tone="teal"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "FFF8E8" if tone == "gold" else PALE_2)
    margins(cell, 90, 110, 90, 110)
    borders(table, color=GOLD if tone == "gold" else TEAL, size="7")
    para = cell.paragraphs[0]
    r = para.add_run(title + "  ")
    set_font(r, size=10.5, bold=True, color=GOLD if tone == "gold" else TEAL)
    r = para.add_run(text)
    set_font(r, size=10.5, italic=True, color=INK)


def add_table(doc, caption, headers, rows, widths, font_size=9.2):
    cap = p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, color=NAVY, before=3, after=2)
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, TEAL)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(header)
        set_font(run, size=font_size, bold=True, color="FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2:
            for cell in cells:
                shade(cell, PALE_2)
        for i, value in enumerate(row):
            cell = cells[i]
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(value))
            set_font(run, size=font_size, color=INK)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Cm(width)
    p(doc, "", after=0)


def add_figure(doc, path, caption, width_cm, height_cm=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run()
    kwargs = {"width": Cm(width_cm)}
    if height_cm:
        kwargs["height"] = Cm(height_cm)
    run.add_picture(str(path), **kwargs)
    cap = p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=MUTED, after=3)
    cap.paragraph_format.keep_with_next = False


def add_figure_pair(doc, left_path, left_caption, right_path, right_caption, width_cm=6.35, height_cm=4.7):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_borders(table)
    for idx, path in enumerate((left_path, right_path)):
        cell = table.cell(0, idx)
        cell.width = Cm(6.8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    for idx, caption in enumerate((left_caption, right_caption)):
        cell = table.cell(1, idx)
        margins(cell, 20, 60, 40, 60)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(caption)
        set_font(run, size=9.5, italic=True, color=MUTED)
    p(doc, "", after=0)


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    p(doc, "WORLD SCIENCE, ENVIRONMENT AND ENGINEERING COMPETITION 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, color=TEAL, after=8)
    rule = doc.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = rule.cell(0, 0)
    cell.width = Cm(14)
    shade(cell, TEAL)
    margins(cell, 18, 0, 18, 0)
    remove_borders(rule)
    p(doc, "FULL PAPER", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, color=NAVY, before=18, after=12)
    p(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, color=INK, after=18)

    logo = doc.add_table(rows=1, cols=1)
    logo.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo.autofit = False
    c = logo.cell(0, 0)
    c.width = Cm(4.8)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(c, PALE_2)
    borders(logo, color=TEAL, size="8")
    margins(c, 430, 100, 430, 100)
    para = c.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("INSERT\nINSTITUTION LOGO")
    set_font(r, size=10.5, bold=True, color=TEAL)

    p(doc, "COMPILED BY", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=TEAL, before=12, after=4)
    for author in AUTHORS:
        p(doc, author, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, color=INK, after=0)

    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    remove_borders(info)
    rows = [
        ("INSTITUTION", "______________________________"),
        ("CATEGORY", "TECHNOLOGY"),
        ("YEAR", "2026"),
    ]
    for ridx, (label, value) in enumerate(rows):
        left, right = info.rows[ridx].cells
        left.width = Cm(4)
        right.width = Cm(8)
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(left.paragraphs[0].add_run(label), size=10, bold=True, color=MUTED)
        set_font(right.paragraphs[0].add_run(value), size=10, bold=True, color=INK)


def add_toc(doc):
    chapter_header(doc, "", "TABLE OF CONTENTS")
    toc_rows = [
        ("Abstract", "3"),
        ("Chapter 1 - Introduction", "4"),
        ("Chapter 2 - Literature Review", "5"),
        ("Chapter 3 - Research Method", "6"),
        ("Chapter 4 - Results and Discussion", "8"),
        ("Chapter 5 - Conclusion", "11"),
        ("References", "12"),
    ]
    p(doc, "This table follows the final Word/PDF pagination of the polished submission version.", align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, italic=True, color=MUTED, after=8)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_borders(table)
    for idx, (label, page) in enumerate(toc_rows):
        cells = table.add_row().cells
        cells[0].width = Cm(12.4)
        cells[1].width = Cm(1.6)
        if idx % 2:
            shade(cells[0], PALE_2)
            shade(cells[1], PALE_2)
        margins(cells[0], 130, 110, 130, 110)
        margins(cells[1], 130, 110, 130, 110)
        set_font(cells[0].paragraphs[0].add_run(label), size=12, bold=True, color=NAVY)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(cells[1].paragraphs[0].add_run(page), size=12, bold=True, color=TEAL)
    callout(doc, "Submission format", "English; A4; Arial 12 pt; single spacing; justified; margins 4/3/3/3 cm; PDF and Microsoft Word.")


def add_abstract(doc):
    chapter_header(doc, "", "ABSTRACT")
    p(doc, ABSTRACT, after=8)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.0
    r = para.add_run("Keywords: ")
    set_font(r, size=12, bold=True, italic=True, color=NAVY)
    r = para.add_run(", ".join(KEYWORDS))
    set_font(r, size=12, italic=True, color=INK)
    callout(doc, "Evidence boundary", "Grade and triage evidence comes from Herlev; KOIL morphology evidence comes from SIPaKMeD conventional Pap-smear crops. Neither is Thai ThinPrep validation or molecular HPV testing.", tone="gold")


def add_chapter1(doc):
    chapter_header(doc, 1, "INTRODUCTION")
    heading(doc, "1.1 Research Background")
    p(doc, "Cervical cancer remains an important public-health problem, yet it can be prevented or treated early when vaccination, screening, diagnosis, and follow-up are connected [1,2]. Cytology-based screening, including Pap and liquid-based preparations such as ThinPrep, examines cellular morphology, while HPV molecular testing detects viral material; the two modalities must not be described as interchangeable [3,4].")
    p(doc, "The practical gap is not limited to image classification. Screening programs also depend on timely review, understandable communication, and reliable referral. In settings with limited cytology expertise, a compact offline decision-support tool may help prioritize abnormal images, expose the image regions that influenced a prediction, and return uncertain cases to human review.")
    heading(doc, "1.2 Problem Statement")
    p(doc, "Many image classifiers provide only a label and confidence score. Such output is insufficient for a safety-sensitive workflow because confidence may be miscalibrated, explanations may be over-trusted, and patient-facing communication may be released without appropriate review. A useful prototype should combine grading, triage, explainability, uncertainty, and explicit clinician control.")
    heading(doc, "1.3 Research Objectives")
    for text in [
        "Develop a four-class Bethesda-style cervical cytology grading model.",
        "Develop a separate koilocytosis-morphology endpoint without treating KOIL as a Bethesda grade.",
        "Add a high-sensitivity binary safety-triage view without replacing the grade output.",
        "Provide Grad-CAM review support and uncertainty-aware abstention.",
        "Demonstrate clinician sign-off, report locking, and offline operation.",
        "Define a validation pathway for Thai ThinPrep images and paired HPV endpoints.",
    ]:
        bullet(doc, text)
    heading(doc, "1.4 Significance")
    p(doc, "The project contributes an integrated medical-AI workflow rather than an isolated classifier. It emphasizes reviewability, conservative communication, reproducible metrics, and explicit boundaries between current evidence and future clinical validation.")


def add_chapter2(doc):
    chapter_header(doc, 2, "LITERATURE REVIEW")
    heading(doc, "2.1 Cervical Cytology and the Bethesda Framework")
    p(doc, "The Bethesda System standardizes cervical cytology terminology and distinguishes negative findings from low-grade, high-grade, and malignant squamous abnormalities [4]. Preserving multiple classes supports clinical communication, whereas a secondary binary view is useful for screening-safety analysis.")
    heading(doc, "2.2 Image Classification and Explainability")
    p(doc, "The Herlev benchmark contains 917 single-cell cervical cytology images [5]. SIPaKMeD provides 4,049 cropped cells from 966 source cluster images across five morphology classes, including 825 koilocytotic cells [14]. EfficientNet provides an efficient scaling strategy for offline image models [6]. Grad-CAM localizes regions associated with a class score [7], but it remains a review aid rather than causal proof.")
    heading(doc, "2.3 Uncertainty, Calibration, and Governance")
    p(doc, "Monte Carlo Dropout provides an approximate uncertainty signal [8], while temperature scaling can improve probability calibration without changing class ordering [9]. Medical-AI guidance emphasizes representative data, human-AI interaction, lifecycle monitoring, and transparent reporting [10-12]. Dataset shift remains a major risk when moving from a public benchmark to a different population, preparation method, microscope, or camera [13].")
    add_table(doc, "Table 1. Design principles derived from the literature.", ["Principle", "System implementation", "Boundary"], [
        ["Structured cytology", "Four-class Bethesda-style grade", "Not a final cytology report"],
        ["KOIL morphology", "Independent one-vs-rest endpoint", "Not HPV DNA/RNA status"],
        ["Screening safety", "Binary normal/abnormal triage", "Not the sole product output"],
        ["Explainability", "Grad-CAM heatmap", "Review aid, not causal proof"],
        ["Uncertainty", "MC Dropout and abstention", "Requires external validation"],
        ["Governance", "Clinician sign-off and report lock", "Demo workflow only"],
    ], [3.1, 6.0, 4.9], font_size=8.7)


def add_chapter3a(doc):
    chapter_header(doc, 3, "RESEARCH METHOD")
    heading(doc, "3.1 Study Design and Dataset")
    p(doc, "This retrospective model-development and prototype-integration study used two public datasets. Grade and binary-triage evidence used 917 real Herlev images [5], with segmentation masks and synthetic experiments excluded. The independent KOIL endpoint used all 4,049 official SIPaKMeD cropped cells from 966 source clusters [14].")
    add_table(doc, "Table 2. Dataset and model card summary.", ["Item", "Specification"], [
        ["Grade/triage dataset", "Herlev: 917 real images; masks excluded"],
        ["KOIL dataset", "SIPaKMeD: 4,049 cells from 966 source clusters"],
        ["Architecture", "EfficientNet-B0 with ImageNet transfer learning"],
        ["Grade output", "NILM, LSIL, HSIL, SCC"],
        ["KOIL output", "Independent koilocytosis morphology probability"],
        ["Evidence status", "Internal public-dataset validation; no Thai ThinPrep evidence"],
    ], [4.1, 9.9], font_size=9.2)
    heading(doc, "3.2 End-to-End Workflow")
    p(doc, "The workflow accepts a cytology image, returns four grade probabilities, binary triage, and an independent KOIL morphology assessment. Grade and KOIL use class-specific Grad-CAM review. Uncertainty and image-quality gates route unsafe cases to clinician confirmation, correction, or rejection; patient-facing output remains locked until all gates pass.")
    add_figure(doc, ASSETS / "workflow_diagram.png", "Figure 1. CerviCo-Pilot clinician-in-the-loop workflow.", 13.2, 4.5)


def add_chapter3b(doc):
    chapter_header(doc, 3, "RESEARCH METHOD", continued=True)
    heading(doc, "3.3 Model Development and Evaluation")
    p(doc, "Images were resized for EfficientNet-B0 input and processed through canonical pipelines. Herlev grade/triage metrics were reproduced from models/best_cervical.pt. For SIPaKMeD, source clusters were split before training (676/145/145 clusters for train/validation/test), preventing cropped cells from the same cluster crossing splits. Model selection, temperature scaling, and a sensitivity-constrained KOIL threshold were fixed on validation data before one-time test evaluation.")
    add_table(doc, "Table 3. Reproducible evaluation specification.", ["Component", "Specification", "Purpose"], [
        ["Held-out evaluation", "n=137; TTA enabled", "Primary test evidence"],
        ["KOIL locked test", "641 cells; 145 clusters", "Independent endpoint evidence"],
        ["Bootstrap", "2,000 percentile resamples", "Confidence intervals"],
        ["Cross-validation", "Five folds; mean +/- SD", "Split robustness"],
        ["Binary threshold", "0.50", "Normal/abnormal triage"],
        ["Calibration", "Temperature=0.794103", "Post-hoc probability scaling"],
        ["KOIL calibration", "Temperature=0.838126", "Validation-fit; test locked"],
        ["Explainability", "Grad-CAM", "Visual review support"],
        ["Uncertainty", "MC Dropout", "Abstention signal"],
    ], [3.5, 5.4, 5.1], font_size=8.5)
    heading(doc, "3.4 Metrics and Safety Policy")
    p(doc, "Legacy grade evaluation included accuracy, macro F1, macro AUROC, per-class recall, and quadratic weighted kappa. Binary triage and KOIL one-vs-rest evaluation included sensitivity, specificity, AUROC, AUPRC, precision, F1, and confusion counts. KOIL confidence intervals used 2,000 source-cluster bootstrap resamples. Calibration used ECE and Brier score. High uncertainty triggers abstention and blocks patient-report release.")
    heading(doc, "3.5 Prototype System")
    p(doc, "The React/FastAPI prototype supports sample review, image upload, separate grade and KOIL probabilities, class-specific Grad-CAM, confirm/edit/reject actions, report export, and an audit demonstration. It is intended for supervised research demonstration and does not satisfy regulated clinical deployment requirements.")
    callout(doc, "Do-not-use boundary", "No autonomous diagnosis, molecular HPV status, unsupervised patient release, or Thai-domain performance claim.", tone="gold")


def add_chapter4a(doc):
    chapter_header(doc, 4, "RESULTS AND DISCUSSION")
    heading(doc, "4.1 Primary Quantitative Results")
    add_table(doc, "Table 4. Held-out and cross-validation results.", ["Measure", "Result", "Interpretation"], [
        ["Legacy five-output accuracy", "0.6934", "KOIL unsupported in Herlev"],
        ["Legacy macro F1 / AUROC", "0.5545 / 0.7311", "Includes unsupported KOIL output"],
        ["Quadratic weighted kappa", "0.687", "Severity agreement"],
        ["Binary sensitivity / specificity", "1.000 / 0.7222", "FN=0; some over-referral"],
        ["Binary AUROC / AUPRC", "0.964 / 0.9856", "Strong Herlev triage"],
        ["Five-fold sensitivity", "0.9867 +/- 0.0086", "High across folds"],
        ["Five-fold AUROC", "0.9435 +/- 0.0448", "Within-domain robustness"],
    ], [5.0, 3.6, 5.4], font_size=8.5)
    add_figure_pair(doc, ASSETS / "headline_metrics.png", "Figure 2. Held-out headline metrics.", ASSETS / "cv_folds.png", "Figure 3. Five-fold cross-validation.", 6.25, 4.55)
    p(doc, "The historical checkpoint used a five-output head, but Herlev contained no true KOIL support; the deployed interface therefore exposes only four grade classes and obtains KOIL morphology from the separate SIPaKMeD model. Binary triage achieved high sensitivity and remains a safety view rather than a replacement for detailed grade review.")


def add_chapter4b(doc):
    chapter_header(doc, 4, "RESULTS AND DISCUSSION", continued=True)
    heading(doc, "4.2 Grade Error Analysis")
    add_figure(doc, KOIL_ASSETS / "herlev_grade_confusion_4class.png", "Figure 4. Herlev confusion matrix for the four supported grade classes.", 7.5, 5.5)
    add_table(doc, "Table 5. Herlev held-out per-class recall.", ["Class", "Support", "Recall", "Interpretation"], [
        ["NILM", "36", "0.7222", "False positives reduce specificity"],
        ["LSIL", "49", "0.6122", "Overlap with higher-grade morphology"],
        ["HSIL", "30", "0.8667", "Strongest high-grade recall"],
        ["SCC", "22", "0.5909", "Requires improvement"],
    ], [2.4, 2.2, 2.2, 7.2], font_size=8.5)
    heading(doc, "4.3 Independent KOIL Morphology Endpoint")
    add_table(doc, "Table 6. Locked SIPaKMeD KOIL one-vs-rest test results.", ["Measure", "Result", "Cluster-bootstrap 95% CI"], [
        ["Support", "133 positive / 508 negative", "145 source clusters"],
        ["Sensitivity / specificity", "0.9624 / 0.9764", "0.9167-0.9921 / 0.9583-0.9916"],
        ["Precision / F1", "0.9143 / 0.9377", "0.8291-0.9702 / 0.8811-0.9706"],
        ["AUROC / AUPRC", "0.9912 / 0.9810", "0.9786-0.9984 / 0.9506-0.9951"],
    ], [4.2, 4.2, 5.6], font_size=8.1)
    p(doc, "At the validation-locked threshold of 0.3367, the model produced 128 TP, 496 TN, 12 FP, and 5 FN cells. False positives were limited to dyskeratotic (n=9) and metaplastic (n=3) morphology; each false negative came from a different source cluster. These errors require expert review and do not represent HPV test errors.", size=10.2)


def add_chapter4c(doc):
    chapter_header(doc, 4, "RESULTS AND DISCUSSION", continued=True)
    heading(doc, "4.4 Calibration, Explainability, and System Interpretation")
    add_table(doc, "Table 7. Held-out calibration evidence.", ["Endpoint / measure", "Before", "After"], [
        ["Multiclass ECE", "0.066853", "0.038670"],
        ["Binary ECE", "0.090679", "0.073787"],
        ["Binary Brier score", "0.071015", "0.066994"],
        ["KOIL ECE", "Not reported", "0.013385"],
        ["KOIL Brier score", "Not reported", "0.019025"],
    ], [6.0, 4.0, 4.0], font_size=9.0)
    add_figure(doc, KOIL_ASSETS / "koil_test_performance.png", "Figure 5. SIPaKMeD morphology confusion matrix and locked-test KOIL ROC/precision-recall curves.", 13.4, 3.6)
    add_figure_pair(doc, ASSETS / "calibration_before_after.png", "Figure 6. Herlev calibration.", KOIL_ASSETS / "koil_gradcam_paper.png", "Figure 7. KOIL-specific Grad-CAM audit cases.", 6.25, 4.25)
    p(doc, "Temperature scaling improved calibration without changing class ordering [9]. KOIL Grad-CAM was generated for the independent endpoint and is presented as attention evidence, not segmentation or causal proof. Calibration and attribution remain domain-dependent and cannot be extended to Thai ThinPrep images without external evidence.", size=10.5)
    p(doc, "The prototype integrates evaluation and review: explanations can be challenged, uncertainty can trigger abstention, and patient communication is gated by clinician action [10-12]. Dataset shift remains the principal translation risk [13].", size=10.5)
    callout(doc, "Required next evidence", "Thai ThinPrep external validation, paired HPV endpoints, pathologist error review, an unaided-versus-AI-assisted reader study, and a prospective workflow pilot.")


def add_chapter5(doc):
    chapter_header(doc, 5, "CONCLUSION")
    heading(doc, "5.1 Conclusion")
    p(doc, "CerviCo-Pilot demonstrates an end-to-end, explainable, and uncertainty-aware cervical cytology screening-support prototype. On Herlev, binary triage achieved held-out sensitivity 1.000 and AUROC 0.964, while the legacy five-output grade evaluation achieved accuracy 0.6934. On the separate locked SIPaKMeD test set, KOIL morphology sensitivity was 0.9624 and AUROC 0.9912. These within-dataset results support continued research but do not justify autonomous diagnosis or deployment.")
    p(doc, "The main contribution is the system design: detailed cytology output is preserved, a binary view emphasizes screening safety, Grad-CAM supports visual review, uncertainty enables abstention, and clinician sign-off controls patient-facing communication. The system can operate locally and presents its limitations alongside its predictions.")
    heading(doc, "5.2 Limitations")
    for text in [
        "Evidence is restricted to two public single-cell conventional cytology datasets.",
        "No Thai ThinPrep/LBC external validation or paired molecular HPV endpoint is available.",
        "KOIL morphology is internally validated on SIPaKMeD only; SCC grade recall remains limited.",
        "Grad-CAM, uncertainty, and report wording have not been evaluated in a clinical reader study.",
        "The web audit and report export remain research-demo components.",
    ]:
        bullet(doc, text)
    heading(doc, "5.3 Future Work")
    p(doc, "The next phase will collect de-identified Thai ThinPrep/LBC images under ethics and data-governance approval, lock an external test set, add paired HPV labels, perform pathologist-reviewed error analysis, and conduct a reader study. A prospective pilot should measure review time, notification time, follow-up completion, and automation bias.")
    heading(doc, "5.4 Research Integrity Statement")
    p(doc, "All performance values were copied from canonical evaluation files for the shipped Herlev grade/triage and SIPaKMeD KOIL checkpoints. External sources support background and methods only; they do not create CerviCo-Pilot performance evidence.")


def add_references(doc):
    chapter_header(doc, "", "REFERENCES")
    for ref in REFERENCES:
        para = p(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.3, color=INK, after=3)
        para.paragraph_format.left_indent = Cm(0.65)
        para.paragraph_format.first_line_indent = Cm(-0.65)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_charts()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.15)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(12)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("CerviCo-Pilot | WSEEC 2026 Full Paper")
    set_font(r, size=8.5, color=MUTED)
    add_page_field(sec.footer.paragraphs[0])

    page_builders = [
        add_cover,
        add_toc,
        add_abstract,
        add_chapter1,
        add_chapter2,
        add_chapter3a,
        add_chapter3b,
        add_chapter4a,
        add_chapter4b,
        add_chapter4c,
        add_chapter5,
        add_references,
    ]
    for idx, builder in enumerate(page_builders):
        builder(doc)
        if idx < len(page_builders) - 1:
            page_break(doc)

    props = doc.core_properties
    props.title = TITLE
    props.subject = "WSEEC 2026 Full Paper - Polished 12-page submission"
    props.author = "CerviCo-Pilot Team"
    props.keywords = ", ".join(KEYWORDS)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
