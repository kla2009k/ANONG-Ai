from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from pythainlp.tokenize import word_tokenize
except Exception:  # pragma: no cover
    word_tokenize = None


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "booklet_assets"
OUT = DOCS / "CerviCo_Pilot_Complete_Report_2026.docx"
SAMPLES = ROOT / "web-react" / "public" / "samples"

COLORS = {
    "teal": "0E7490",
    "navy": "143B5E",
    "ink": "0F2A3A",
    "muted": "5B7282",
    "line": "DDE7EC",
    "soft": "EEF7FA",
    "gold": "C98A1E",
    "red": "C0492F",
    "green": "2F8F6B",
    "koil": "6D5BD0",
}


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def thai_break(text: str) -> str:
    if not text or word_tokenize is None:
        return text
    # Preserve explicit spaces; add ZWSP only inside Thai-heavy spans.
    out = []
    for part in text.split(" "):
        if any("\u0e00" <= ch <= "\u0e7f" for ch in part) and len(part) > 12:
            try:
                out.append("\u200b".join(t for t in word_tokenize(part, engine="newmm") if t))
            except Exception:
                out.append(part)
        else:
            out.append(part)
    return " ".join(out)


def set_run_font(run, name="TH Sarabun New", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DDE7EC", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(row.cells[idx])


def style_doc(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "TH Sarabun New")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "TH Sarabun New")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    normal._element.rPr.rFonts.set(qn("w:cs"), "TH Sarabun New")
    normal.font.size = Pt(16)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 20, COLORS["navy"], 14, 8),
        ("Heading 2", 18, COLORS["teal"], 10, 6),
        ("Heading 3", 16, COLORS["ink"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = "TH Sarabun New"
        st._element.rPr.rFonts.set(qn("w:ascii"), "TH Sarabun New")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "TH Sarabun New")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        st._element.rPr.rFonts.set(qn("w:cs"), "TH Sarabun New")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("CerviCo-Pilot Complete Report 2026")
    set_run_font(r, size=11, color=COLORS["muted"])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Phase 1 Herlev-only research prototype | Decision-support only")
    set_run_font(r, size=10, color=COLORS["muted"])


def add_p(doc, text="", style=None, align=None, bold=False, color=None, size=16):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(thai_break(text))
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    return add_p(doc, text, style=f"Heading {level}")


def add_callout(doc, title, body, tone="soft"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [15.7])
    set_table_borders(table, color=COLORS["line"], size="6")
    cell = table.cell(0, 0)
    shade_cell(cell, COLORS["soft"] if tone == "soft" else "FFF7E6")
    p = cell.paragraphs[0]
    r = p.add_run(thai_break(title))
    set_run_font(r, size=16, bold=True, color=COLORS["navy"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(thai_break(body))
    set_run_font(r2, size=15, color=COLORS["ink"])


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(thai_break(item))
        set_run_font(r, size=16)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(thai_break(item))
        set_run_font(r, size=16)


def add_table(doc, caption, headers, rows, widths_cm=None, source=None):
    add_p(doc, caption, bold=True, color=COLORS["navy"])
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], "E8EEF5")
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(thai_break(str(h)))
        set_run_font(r, size=14, bold=True, color=COLORS["navy"])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) < 22 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(thai_break(str(val)))
            set_run_font(r, size=13)
    if widths_cm:
        set_table_width(table, widths_cm)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    if source:
        add_p(doc, f"ที่มา: {source}", size=12, color=COLORS["muted"])
    return table


def add_figure(doc, image_path: Path, caption: str, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(thai_break(caption))
    set_run_font(r, size=13, color=COLORS["muted"])


def make_charts():
    ASSETS.mkdir(parents=True, exist_ok=True)
    tri = load_json(ROOT / "models" / "triage_metrics.json")
    cv = load_json(ROOT / "models" / "cv_results.json")
    curves = load_json(SAMPLES / "curves.json")
    cal = load_json(ROOT / "models" / "calibration_temperature.json")

    plt.style.use("seaborn-v0_8-whitegrid")

    # Headline metrics
    fig, ax = plt.subplots(figsize=(8, 4.8), dpi=180)
    labels = ["5-class acc", "Macro AUROC", "HSIL recall", "SCC recall", "Binary sens", "Binary AUROC"]
    vals = [
        tri["five_class"]["accuracy"],
        tri["five_class"]["auc_macro"],
        tri["five_class"]["per_class"]["HSIL"]["recall"],
        tri["five_class"]["per_class"]["SCC"]["recall"],
        tri["binary_triage"]["sensitivity_recall"],
        tri["binary_triage"]["auroc"],
    ]
    colors = ["#0E7490", "#143B5E", "#C98A1E", "#C0492F", "#2F8F6B", "#0E7490"]
    ax.bar(labels, vals, color=colors)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Score")
    ax.set_title("CerviCo-Pilot Phase 1 Held-out Metrics")
    ax.tick_params(axis="x", rotation=25)
    for i, v in enumerate(vals):
        ax.text(i, v + 0.025, f"{v:.3f}", ha="center", fontsize=8)
    fig.tight_layout()
    fig.savefig(ASSETS / "headline_metrics.png", bbox_inches="tight")
    plt.close(fig)

    # CV folds
    folds = cv["per_fold"]
    x = np.arange(1, len(folds) + 1)
    fig, ax = plt.subplots(figsize=(8, 4.8), dpi=180)
    ax.plot(x, [f["acc5"] for f in folds], marker="o", label="5-class acc")
    ax.plot(x, [f["tri_sens"] for f in folds], marker="o", label="Binary sens")
    ax.plot(x, [f["tri_auroc"] for f in folds], marker="o", label="Binary AUROC")
    ax.set_ylim(0.45, 1.04)
    ax.set_xlabel("Fold")
    ax.set_ylabel("Score")
    ax.set_title("5-fold Cross-validation")
    ax.legend(loc="lower right")
    fig.tight_layout()
    fig.savefig(ASSETS / "cv_folds.png", bbox_inches="tight")
    plt.close(fig)

    # Confusion matrix
    matrix = np.array([[26, 1, 7, 2, 0], [1, 30, 16, 2, 0], [1, 0, 26, 3, 0], [0, 0, 9, 13, 0], [0, 0, 0, 0, 0]])
    classes = ["NILM", "LSIL", "HSIL", "SCC", "KOIL"]
    fig, ax = plt.subplots(figsize=(6.2, 5.3), dpi=180)
    display_matrix = matrix.astype(float)
    display_matrix[-1, :] = np.nan
    cmap = plt.cm.Blues.copy()
    cmap.set_bad("#ECEFF1")
    im = ax.imshow(np.ma.masked_invalid(display_matrix), cmap=cmap)
    ax.set_xticks(range(len(classes)), labels=classes)
    ax.set_yticks(range(len(classes)), labels=classes)
    ax.set_xlabel("Predicted")
    ax.set_ylabel("True")
    ax.set_title("Held-out 5-class Confusion Matrix\nKOIL recall not estimable (support = 0)")
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            if i < matrix.shape[0] - 1:
                ax.text(j, i, str(matrix[i, j]), ha="center", va="center", color="#0F2A3A", fontsize=9)
    ax.text(2, 4, "N/A - no true KOIL samples", ha="center", va="center", color="#526773", fontsize=8.5, fontweight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(ASSETS / "confusion_matrix.png", bbox_inches="tight")
    plt.close(fig)

    # ROC
    fig, ax = plt.subplots(figsize=(6.5, 5), dpi=180)
    ax.plot([p["fpr"] for p in curves["roc"]], [p["tpr"] for p in curves["roc"]], color="#0E7490", linewidth=2.5)
    ax.plot([0, 1], [0, 1], "--", color="#999999")
    ax.set_xlabel("False positive rate")
    ax.set_ylabel("True positive rate")
    ax.set_title(f"Binary Triage ROC (AUROC {curves['roc_auc']:.3f})")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    fig.savefig(ASSETS / "roc_curve.png", bbox_inches="tight")
    plt.close(fig)

    # Calibration before/after
    fig, ax = plt.subplots(figsize=(7.2, 4.8), dpi=180)
    labels = ["Multiclass ECE", "Binary ECE", "Binary Brier"]
    before = [
        cal["test_before"]["ece_multiclass_10bin"],
        cal["test_before"]["ece_binary_10bin"],
        cal["test_before"]["brier_binary"],
    ]
    after = [
        cal["test_after"]["ece_multiclass_10bin"],
        cal["test_after"]["ece_binary_10bin"],
        cal["test_after"]["brier_binary"],
    ]
    idx = np.arange(len(labels))
    ax.bar(idx - 0.18, before, 0.36, label="Before", color="#C98A1E")
    ax.bar(idx + 0.18, after, 0.36, label="After temp scaling", color="#0E7490")
    ax.set_xticks(idx, labels)
    ax.set_ylabel("Lower is better")
    ax.set_title("Post-hoc Temperature Scaling on Held-out Herlev")
    ax.legend()
    fig.tight_layout()
    fig.savefig(ASSETS / "calibration_before_after.png", bbox_inches="tight")
    plt.close(fig)

    # Sample grid
    pairs = []
    for i, label in enumerate(["NILM", "NILM", "LSIL", "LSIL", "HSIL", "HSIL", "SCC", "SCC"]):
        pairs.append((SAMPLES / f"s{i:02d}_{label}.jpg", SAMPLES / f"s{i:02d}_{label}_cam.jpg", f"s{i:02d} {label}"))
    thumb_w, thumb_h = 180, 180
    canvas = Image.new("RGB", (thumb_w * 4, (thumb_h + 28) * 4), "white")
    draw = ImageDraw.Draw(canvas)
    for idx, (orig, cam, label) in enumerate(pairs[:8]):
        row = (idx // 2)
        col = (idx % 2) * 2
        for k, img_path in enumerate([orig, cam]):
            img = Image.open(img_path).convert("RGB")
            img.thumbnail((thumb_w, thumb_h))
            x = (col + k) * thumb_w + (thumb_w - img.width) // 2
            y = row * (thumb_h + 28)
            canvas.paste(img, (x, y))
        draw.text((col * thumb_w + 4, row * (thumb_h + 28) + thumb_h + 4), f"{label}: original / Grad-CAM", fill=(15, 42, 58))
    canvas.save(ASSETS / "sample_gradcam_grid.jpg", quality=92)

    # Workflow diagram
    fig, ax = plt.subplots(figsize=(9, 4.6), dpi=180)
    ax.axis("off")
    steps = ["Image", "5-class\nBethesda", "HPV morphology\nrisk", "Grad-CAM +\nuncertainty", "Clinician\nsign-off", "Report +\nfollow-up"]
    xs = np.linspace(0.08, 0.92, len(steps))
    for i, (xpos, step) in enumerate(zip(xs, steps)):
        rect = plt.Rectangle((xpos - 0.065, 0.43), 0.13, 0.23, facecolor="#EEF7FA", edgecolor="#0E7490", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(xpos, 0.545, step, ha="center", va="center", fontsize=9, color="#0F2A3A")
        if i < len(steps) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.075, 0.545), xytext=(xpos + 0.075, 0.545), arrowprops=dict(arrowstyle="->", color="#143B5E", lw=1.8))
    ax.text(0.5, 0.23, "Safety: high uncertainty blocks patient report; AI does not replace HPV DNA/RNA testing or clinician review.", ha="center", fontsize=9, color="#C0492F")
    fig.tight_layout()
    fig.savefig(ASSETS / "workflow_diagram.png", bbox_inches="tight")
    plt.close(fig)


def build_doc():
    tri = load_json(ROOT / "models" / "triage_metrics.json")
    cv = load_json(ROOT / "models" / "cv_results.json")
    cal = load_json(ROOT / "models" / "calibration_temperature.json")

    make_charts()

    doc = Document()
    style_doc(doc)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("รายงานฉบับสมบูรณ์")
    set_run_font(r, size=30, bold=True, color=COLORS["navy"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("การพัฒนาระบบปัญญาประดิษฐ์เพื่อคัดกรองความผิดปกติของเซลล์ปากมดลูก\nและการประเมินความเสี่ยงเชื้อ HPV จากภาพถ่าย ThinPrep")
    set_run_font(r, size=24, bold=True, color=COLORS["teal"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Development of an AI System for Cervical Cell Abnormality Screening and HPV Risk Assessment from ThinPrep Images")
    set_run_font(r, size=16, color=COLORS["muted"])
    add_figure(doc, ASSETS / "workflow_diagram.png", "ภาพรวมระบบ CerviCo-Pilot: workflow ตั้งแต่ภาพเซลล์ถึงการลงนามโดยแพทย์", width_cm=15.5)
    add_p(doc, "CerviCo-Pilot | Phase 1 Herlev-only Research Prototype | 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=COLORS["ink"], bold=True)
    add_callout(doc, "ขอบเขตสำคัญ", "รายงานนี้อธิบาย prototype เพื่อการวิจัย/ประกวด ไม่ใช่เอกสารรับรองการใช้งานทางคลินิก ระบบไม่วินิจฉัยแทนแพทย์ และไม่ตรวจหา HPV DNA/RNA", "warn")
    doc.add_page_break()

    # Abstracts
    add_heading(doc, "บทคัดย่อ", 1)
    add_p(doc, "CerviCo-Pilot เป็นระบบช่วยคัดกรองภาพเซลล์วิทยาปากมดลูกแบบ clinician-in-the-loop โดยอ่านภาพ Pap/ThinPrep-style และให้ผลลัพธ์แบบ Bethesda-style 5 class พร้อม binary safety triage, Grad-CAM, uncertainty flag, HPV-related morphology risk และการล็อกรายงานผู้ป่วยจนกว่าบุคลากรแพทย์จะยืนยันผล")
    add_p(doc, "โมเดลปัจจุบันเป็น EfficientNet-B0 ที่ประเมินบนภาพจริงจากชุดข้อมูลสาธารณะ Herlev จำนวน 917 ภาพ โดยตัด segmentation mask ออกและใช้ผลจาก checkpoint ปัจจุบันเท่านั้น ผล held-out 5-class accuracy เท่ากับ 0.6934 และ macro AUROC เท่ากับ 0.7311 ส่วน binary triage มี sensitivity 1.0 และ AUROC 0.964 ใน held-out test ขณะที่ 5-fold binary sensitivity เท่ากับ 0.9867 +/- 0.0086")
    add_p(doc, "จุดสำคัญของโครงงานไม่ใช่การแทนที่ผู้เชี่ยวชาญ แต่คือการสร้าง workflow ที่ช่วยให้ความผิดปกติถูกสังเกต อธิบาย และส่งต่อได้ปลอดภัยขึ้น ระบบยังไม่มี Thai ThinPrep validation, paired HPV DNA/RNA endpoint, reader study หรือ prospective workflow trial จึงต้องสื่อสารเป็น Phase 1 research prototype เท่านั้น")
    add_heading(doc, "Abstract", 1)
    add_p(doc, "CerviCo-Pilot is a clinician-in-the-loop cervical cytology screening support prototype. It reads Pap/ThinPrep-style cytology images and returns a Bethesda-style 5-class suggestion, binary safety-triage view, Grad-CAM heatmap, uncertainty signal, and HPV-related morphology risk note. Patient-facing reports remain locked until clinician sign-off.")
    add_p(doc, "The current model is an EfficientNet-B0 evaluated on 917 real public Herlev images with segmentation masks excluded. Held-out 5-class accuracy is 0.6934 and macro AUROC is 0.7311. The binary triage layer achieves held-out sensitivity 1.0 and AUROC 0.964, with 5-fold sensitivity 0.9867 +/- 0.0086. These findings support a Phase 1 research prototype, not clinical deployment.")
    doc.add_page_break()

    # Manual TOC
    add_heading(doc, "สารบัญ", 1)
    toc_rows = [
        ["บทที่ 1", "บทนำและปัญหา", "แนวคิด โจทย์ วัตถุประสงค์ และขอบเขต"],
        ["บทที่ 2", "หลักฐานภายนอกและกรอบทางคลินิก", "WHO, CDC, FDA/IMDRF, CONSORT-AI, DECIDE-AI"],
        ["บทที่ 3", "วิธีดำเนินงานและสถาปัตยกรรมระบบ", "dataset, model, XAI, uncertainty, UX workflow"],
        ["บทที่ 4", "ผลการทดลองและการวิเคราะห์", "metrics, ROC, CV, confusion matrix, calibration, case gallery"],
        ["บทที่ 5", "การนำเสนอ การใช้งานเดโม และแผนพัฒนา", "submission, demo runbook, risks, roadmap"],
        ["ภาคผนวก", "เอกสารกำกับ", "intended use, Q&A, failure modes, source ledger"],
    ]
    add_table(doc, "ตารางที่ 0.1 โครงสร้างรูปเล่ม", ["ส่วน", "หัวข้อ", "รายละเอียด"], toc_rows, [2.5, 5.0, 8.2])
    add_heading(doc, "รายการรูป", 2)
    add_bullets(doc, [
        "รูปที่ 1: workflow diagram ของ CerviCo-Pilot",
        "รูปที่ 2: กราฟ headline metrics",
        "รูปที่ 3: 5-fold cross-validation",
        "รูปที่ 4: confusion matrix",
        "รูปที่ 5: ROC curve ของ binary triage",
        "รูปที่ 6: calibration ก่อน/หลัง temperature scaling",
        "รูปที่ 7: sample cytology + Grad-CAM grid",
    ])
    doc.add_page_break()

    # Chapter 1
    add_heading(doc, "บทที่ 1 บทนำและปัญหา", 1)
    add_heading(doc, "1.1 ที่มาและความสำคัญ", 2)
    add_p(doc, "มะเร็งปากมดลูกเป็นปัญหาสาธารณสุขที่สามารถลดความรุนแรงได้เมื่อระบบคัดกรอง การสื่อสารผล และการติดตามรักษาทำงานร่วมกันอย่างต่อเนื่อง โครงงานนี้จึงไม่ได้มอง AI เป็นเครื่องมือวินิจฉัยอัตโนมัติ แต่เป็นผู้ช่วยใน workflow ที่ทำให้ภาพเซลล์ผิดปกติถูกจัดลำดับ อธิบาย และส่งต่อให้บุคลากรแพทย์ตรวจทานได้เร็วขึ้น")
    add_p(doc, "ในบริบทการประกวดและงานวิจัยระดับต้น สิ่งสำคัญคือความซื่อสัตย์ของ claim ระบบจึงระบุชัดว่าเป็น Phase 1 prototype ที่มีหลักฐานจาก Herlev public dataset เท่านั้น ยังไม่มีข้อมูล Thai ThinPrep/LBC และยังไม่ใช่ clinical deployment")
    add_heading(doc, "1.2 วัตถุประสงค์", 2)
    add_numbered(doc, [
        "พัฒนาระบบช่วยคัดกรองภาพเซลล์ปากมดลูกแบบ Bethesda-style 5-class",
        "เพิ่ม binary safety triage เพื่อเน้นการลด false negative ในระดับ screening",
        "แสดง Grad-CAM เพื่อช่วยให้แพทย์ตรวจทานบริเวณที่โมเดลให้ความสำคัญ",
        "เพิ่ม uncertainty gate เพื่อให้ระบบหยุดและส่งคืนให้มนุษย์เมื่อไม่มั่นใจ",
        "สร้างรายงานแพทย์และรายงานผู้ป่วยที่ปล่อยได้หลัง clinician sign-off เท่านั้น",
        "วางกรอบ HPV-related morphology risk โดยไม่อ้างว่าเป็น HPV DNA/RNA test",
    ])
    add_heading(doc, "1.3 ขอบเขตและข้อจำกัด", 2)
    add_table(
        doc,
        "ตารางที่ 1.1 ขอบเขตการใช้งานและข้อห้าม",
        ["รายการ", "สถานะในโครงงาน", "คำอธิบายที่ปลอดภัย"],
        [
            ["ชนิดระบบ", "Decision-support prototype", "ช่วยคัดกรองและจัดลำดับ ไม่ใช่เครื่องมือวินิจฉัยสุดท้าย"],
            ["ข้อมูล", "Herlev public dataset", "ยังไม่มี Thai ThinPrep validation"],
            ["HPV", "Morphology risk only", "ไม่ตรวจหาเชื้อหรือสารพันธุกรรม HPV"],
            ["KOIL", "Phase 2 target", "Herlev ไม่มี true KOIL จึงไม่ใช่ capability ปัจจุบัน"],
            ["Audit", "Local demo", "localStorage ใช้สาธิต ต้องเปลี่ยนเป็น server-side signed log ใน pilot"],
        ],
        [3.2, 4.0, 8.5],
        "docs/INTENDED_USE_STATEMENT.md และ docs/CLAIMS_LEDGER.md",
    )

    # Chapter 2
    add_heading(doc, "บทที่ 2 หลักฐานภายนอกและกรอบทางคลินิก", 1)
    add_heading(doc, "2.1 หลักฐานด้านสาธารณสุข", 2)
    add_p(doc, "WHO ระบุว่ามะเร็งปากมดลูกสามารถป้องกันได้เป็นส่วนใหญ่ผ่าน HPV vaccination, screening และ treatment/follow-up pathways ขณะที่กลยุทธ์ elimination ใช้เป้าหมาย 90-70-90 ซึ่งทำให้โครงงานนี้ควรผูกกับระบบคัดกรองและการติดตามผล ไม่ใช่การกล่าวว่า AI แก้ปัญหาทั้งหมดด้วยตัวเอง")
    add_heading(doc, "2.2 ความแตกต่างระหว่าง cytology และ HPV testing", 2)
    add_p(doc, "CDC แยกชัดเจนว่า HPV test ตรวจหาไวรัส ส่วน Pap test ตรวจหาความเปลี่ยนแปลงของเซลล์ ดังนั้น CerviCo-Pilot จึงสามารถพูดเรื่อง HPV ได้เฉพาะในความหมายของ HPV-related morphology risk จากภาพ ไม่ใช่การตรวจยืนยัน infection")
    add_table(
        doc,
        "ตารางที่ 2.1 แหล่งอ้างอิงภายนอกและการใช้งานในรายงาน",
        ["แหล่ง", "ใช้สนับสนุน", "ข้อจำกัด"],
        [
            ["WHO cervical cancer fact sheet", "ปัญหา prevention, screening, treatment/follow-up", "ไม่ใช่หลักฐาน performance ของโมเดล"],
            ["CDC cervical screening", "แยก Pap/cytology ออกจาก HPV test", "ใช้คุม wording ไม่ให้กล่าวเกินจริง"],
            ["FDA/IMDRF GMLP", "lifecycle governance ของ medical AI", "เป็นหลักการ ไม่ใช่ regulatory clearance"],
            ["CONSORT-AI", "human-AI interaction และ error analysis", "ใช้วาง roadmap ไม่ใช่ผลทดลองปัจจุบัน"],
            ["DECIDE-AI", "early-stage clinical evaluation", "ชี้ว่าต้องมี reader study ก่อน clinical utility claim"],
        ],
        [4.0, 6.3, 5.4],
        "docs/SOURCE_CITATION_LEDGER.md",
    )
    add_callout(doc, "คำตอบหลักเรื่อง HPV", "ภาพ cytology ไม่ได้ตรวจหาเชื้อ HPV โดยตรง ระบบนี้ประเมินความเสี่ยงที่สัมพันธ์กับ HPV จาก morphology ของเซลล์ และต้องแยกจาก HPV DNA/RNA testing เสมอ", "warn")

    # Chapter 3
    add_heading(doc, "บทที่ 3 วิธีดำเนินงานและสถาปัตยกรรมระบบ", 1)
    add_figure(doc, ASSETS / "workflow_diagram.png", "รูปที่ 3.1 สถาปัตยกรรม workflow ของ CerviCo-Pilot", 15.5)
    add_heading(doc, "3.1 Dataset และการจัดเตรียมข้อมูล", 2)
    add_p(doc, "ข้อมูลหลักในระยะ Phase 1 คือ Herlev public cervical cytology dataset จำนวน 917 ภาพ โดยใช้เฉพาะภาพจริงและตัด segmentation mask ออก เพื่อป้องกันการปนเปื้อนของหลักฐานจาก mask หรือข้อมูล synthetic-heavy run เก่า")
    add_table(
        doc,
        "ตารางที่ 3.1 Dataset/model card แบบย่อ",
        ["หัวข้อ", "รายละเอียด"],
        [
            ["Dataset", "Herlev public cervical cytology"],
            ["จำนวนภาพ", "917 real images"],
            ["Model", "EfficientNet-B0"],
            ["Output", "NILM, LSIL, HSIL, SCC, KOIL placeholder"],
            ["Evidence level", "Phase 1 retrospective public-dataset prototype"],
            ["Not yet done", "Thai ThinPrep validation, paired HPV endpoint, reader study, prospective pilot"],
        ],
        [4.5, 11.0],
        "docs/DATASET_MODEL_CARD.md",
    )
    add_heading(doc, "3.2 โมเดลและ output layers", 2)
    add_bullets(doc, [
        "5-class Bethesda-style output เป็น product identity เพราะ cytology workflow ต้องการรายละเอียดมากกว่า normal/abnormal",
        "Binary normal/abnormal triage เป็น safety layer ที่ช่วยอธิบาย sensitivity ในบริบท screening",
        "HPV-related morphology risk เป็น clinical explanation layer ไม่ใช่ molecular endpoint",
        "Grad-CAM เป็น review aid ให้แพทย์ตรวจทาน ไม่ใช่ causal explanation",
        "MC Dropout/uncertainty เป็น abstention signal สำหรับบล็อกรายงานผู้ป่วยเมื่อระบบไม่มั่นใจ",
    ])
    add_heading(doc, "3.3 Web demo และ clinical workflow", 2)
    add_table(
        doc,
        "ตารางที่ 3.2 หน้าเว็บและหน้าที่ของแต่ละหน้า",
        ["Route", "หน้าที่", "เหตุผลเชิงกรรมการ"],
        [
            ["/", "Dashboard current truth", "เห็นสถานะจริงทันที"],
            ["/analyze", "วิเคราะห์ภาพ/ตัวอย่าง + sign-off", "โชว์ workflow หลัก"],
            ["/gallery", "case gallery รวมเคสถูกและพลาด", "เพิ่มความน่าเชื่อถือ ไม่ cherry-pick"],
            ["/workflow", "clinical pathway + safety gates", "อธิบายว่าระบบไม่ใช่ classifier ลอย ๆ"],
            ["/reports", "รายงานแพทย์/ผู้ป่วยและ lock state", "แสดง patient safety"],
            ["/history", "local audit trail", "พิสูจน์แนวคิด governance"],
            ["/model", "model card/do-not-use", "คุม claim และ intended use"],
        ],
        [2.7, 5.2, 7.6],
        "web-react/src/App.tsx",
    )

    # Chapter 4
    add_heading(doc, "บทที่ 4 ผลการทดลองและการวิเคราะห์", 1)
    add_heading(doc, "4.1 Headline metrics", 2)
    add_table(
        doc,
        "ตารางที่ 4.1 ผล held-out และ cross-validation ที่ใช้กล่าวอ้างได้",
        ["Metric", "Value", "Interpretation"],
        [
            ["Held-out 5-class accuracy", "0.6934", "moderate 5-class grading"],
            ["Held-out macro AUROC", "0.7311", "moderate multiclass discrimination"],
            ["Held-out HSIL recall", "0.8667", "strong HSIL recall"],
            ["Held-out SCC recall", "0.5909", "SCC remains imperfect"],
            ["Held-out binary sensitivity", "1.0", "no false negatives in held-out triage test"],
            ["Held-out binary AUROC", "0.964", "strong binary discrimination"],
            ["5-fold binary sensitivity", "0.9867 +/- 0.0086", "robust high sensitivity"],
            ["5-fold binary AUROC", "0.9435 +/- 0.0448", "robust triage AUROC"],
        ],
        [4.6, 3.5, 7.5],
        "models/test_metrics.json, models/triage_metrics.json, models/cv_results.json",
    )
    add_figure(doc, ASSETS / "headline_metrics.png", "รูปที่ 4.1 กราฟสรุป metrics สำคัญจาก held-out test", 15.0)
    add_figure(doc, ASSETS / "cv_folds.png", "รูปที่ 4.2 ผล 5-fold cross-validation: 5-class accuracy, binary sensitivity และ binary AUROC", 15.0)
    add_heading(doc, "4.2 Confusion matrix และข้อจำกัดราย class", 2)
    add_figure(doc, ASSETS / "confusion_matrix.png", "รูปที่ 4.3 Confusion matrix แบบ 5-class บน held-out test", 12.0)
    add_p(doc, "ผล 5-class มีข้อจำกัดสำคัญ: HSIL recall สูงกว่า SCC recall ขณะที่ KOIL recall ไม่สามารถประมาณได้ (N/A) เพราะไม่มี true KOIL support ใน Herlev จึงต้องสื่อสาร KOIL เป็น Phase 2 target เท่านั้น")
    add_heading(doc, "4.3 ROC และ calibration", 2)
    add_figure(doc, ASSETS / "roc_curve.png", "รูปที่ 4.4 ROC curve ของ binary triage layer", 12.5)
    add_figure(doc, ASSETS / "calibration_before_after.png", "รูปที่ 4.5 ผล post-hoc temperature scaling บน Herlev held-out test", 14.0)
    add_table(
        doc,
        "ตารางที่ 4.2 Calibration ก่อนและหลัง temperature scaling บน test split",
        ["Metric", "Before", "After", "Note"],
        [
            ["Temperature", "1.000000", f"{cal['temperature']:.6f}", "scalar post-hoc temperature"],
            ["Multiclass ECE", f"{cal['test_before']['ece_multiclass_10bin']:.6f}", f"{cal['test_after']['ece_multiclass_10bin']:.6f}", "improved on Herlev"],
            ["Binary ECE", f"{cal['test_before']['ece_binary_10bin']:.6f}", f"{cal['test_after']['ece_binary_10bin']:.6f}", "improved on Herlev"],
            ["Binary Brier", f"{cal['test_before']['brier_binary']:.6f}", f"{cal['test_after']['brier_binary']:.6f}", "slightly improved"],
            ["Binary AUROC", f"{cal['test_before']['auroc_binary']:.6f}", f"{cal['test_after']['auroc_binary']:.6f}", "nearly unchanged"],
        ],
        [4.2, 3.0, 3.0, 5.2],
        "models/calibration_temperature.json",
    )
    add_callout(doc, "ข้อควรระวังเรื่อง calibration", "Temperature scaling ทำให้ ECE/Brier ดีขึ้นบน Herlev held-out test แต่ยังไม่ใช่ external Thai ThinPrep calibration และไม่ควรกล่าวว่า fully calibrated", "warn")
    add_heading(doc, "4.4 รูปตัวอย่างและ Grad-CAM", 2)
    add_figure(doc, ASSETS / "sample_gradcam_grid.jpg", "รูปที่ 4.6 ตัวอย่างภาพ cytology และ Grad-CAM จาก Herlev samples", 15.5)

    # Chapter 5
    add_heading(doc, "บทที่ 5 การใช้งานเดโม การส่งประกวด และแผนพัฒนา", 1)
    add_heading(doc, "5.1 เส้นทางการสาธิตเว็บ", 2)
    add_numbered(doc, [
        "เปิด dashboard เพื่อย้ำว่าเป็น Phase 1 Herlev-only prototype",
        "เปิดหน้า Analyze เลือก sample และแสดง 5-class result, probability, HPV morphology note และ Grad-CAM",
        "กด clinician confirm/edit/reject เพื่อแสดง human-in-the-loop",
        "เปิด History เพื่อแสดง local audit trail",
        "เปิด Gallery เพื่อแสดงทั้งเคสถูกและเคสพลาด",
        "เปิด Reports เพื่อแสดง patient-report lock/unlock behavior",
        "เปิด Model Card เพื่อย้ำ do-not-use boundaries",
    ])
    add_heading(doc, "5.2 การวาง narrative สำหรับเวที", 2)
    add_table(
        doc,
        "ตารางที่ 5.1 ความต่างของ narrative ระหว่าง Samsung SFT และ WSEEC",
        ["เวที", "ควรเน้น", "ไม่ควรนำด้วย"],
        [
            ["Samsung Solve for Tomorrow", "social impact, follow-up workflow, patient communication, offline prototype", "technical metrics หนาแน่นเกินไปหรือ claim clinical readiness"],
            ["WSEEC", "scientific method, Herlev dataset, metrics, calibration, error analysis, validation roadmap", "social impact โดยไม่มี evidence และตัวเลข unsupported"],
        ],
        [4.0, 6.6, 5.1],
        "docs/SFT_WSEEC_SUBMISSION_PACKAGE.md",
    )
    add_heading(doc, "5.3 แผนพัฒนา", 2)
    add_table(
        doc,
        "ตารางที่ 5.2 Roadmap เพื่อขยับจาก prototype ไปสู่ pilot",
        ["ระยะ", "งานหลัก", "Acceptance criteria"],
        [
            ["Phase 2 data", "Thai ThinPrep/LBC retrospective dataset", "locked Thai test set, no leakage, per-class support"],
            ["HPV endpoint", "paired HPV DNA/RNA labels", "AUROC/AUPRC and calibration for HPV-positive endpoint"],
            ["Error review", "pathologist-reviewed gallery", "case-level notes and Grad-CAM failure review"],
            ["Reader study", "unaided vs AI-assisted", "sensitivity, specificity, time, overreliance analysis"],
            ["Pilot governance", "server-side signed audit", "user identity, model version, immutable log"],
            ["Prospective workflow", "ethics-approved pilot", "follow-up completion/time-to-notification outcomes"],
        ],
        [3.1, 5.8, 6.8],
        "docs/VALIDATION_ROADMAP.md และ docs/SERVER_SIDE_AUDIT_ROADMAP.md",
    )

    # Appendices
    doc.add_page_break()
    add_heading(doc, "ภาคผนวก ก Intended Use แบบย่อ", 1)
    add_table(
        doc,
        "ตารางภาคผนวก ก.1 Intended use summary",
        ["หัวข้อ", "คำตอบ"],
        [
            ["Intended use", "Clinician-in-the-loop cervical cytology screening support prototype"],
            ["Intended users", "cytotechnologists, pathologists, trained clinicians, supervised evaluators"],
            ["Inputs", "Pap/ThinPrep-style cytology images; bundled Herlev samples; future Thai metadata"],
            ["Outputs", "5-class suggestion, binary triage, HPV morphology note, Grad-CAM, uncertainty, reports"],
            ["Do not use", "final diagnosis, HPV infection detection, autonomous patient release, Thai validation claim"],
        ],
        [4.0, 11.7],
        "docs/INTENDED_USE_STATEMENT.md",
    )
    add_heading(doc, "ภาคผนวก ข Judge Q&A แบบย่อ", 1)
    qa_rows = [
        ["Why not ChatGPT?", "General LLMs can explain confirmed results, but they are not the evaluated cytology classifier."],
        ["Why not HPV DNA test?", "HPV testing detects virus; CerviCo-Pilot supports cytology interpretation and follow-up workflow."],
        ["Is it clinically validated?", "No. It is Phase 1 Herlev-only evidence; Thai validation and reader study remain future work."],
        ["Why keep 5 classes?", "5-class is clinical detail/product identity; binary triage is the safety layer."],
        ["What if uncertain?", "High uncertainty blocks patient report and routes the case back to human review."],
    ]
    add_table(doc, "ตารางภาคผนวก ข.1 คำถามกรรมการที่ควรเตรียมตอบ", ["คำถาม", "คำตอบสั้นที่ปลอดภัย"], qa_rows, [4.6, 11.1], "docs/JUDGE_QA_BANK.md")
    add_heading(doc, "ภาคผนวก ค Failure modes", 1)
    add_table(
        doc,
        "ตารางภาคผนวก ค.1 Failure modes และ mitigation",
        ["Failure mode", "Current response", "Future control"],
        [
            ["Automation bias", "UI shows uncertainty, probabilities, Grad-CAM, sign-off", "reader study for overreliance"],
            ["Patient misunderstanding", "patient report lock and disclaimers", "Thai health-literacy testing"],
            ["Grad-CAM overtrust", "gallery includes wrong predictions", "pathologist-reviewed heatmap failure gallery"],
            ["Out-of-domain image", "no Thai validation claim", "Thai locked external validation"],
            ["Local audit overclaim", "documented as demo-only", "server-side signed audit"],
        ],
        [4.3, 5.8, 5.6],
        "docs/FAILURE_MODE_AND_HUMAN_FACTORS.md",
    )
    add_heading(doc, "ภาคผนวก ง Source ledger แบบย่อ", 1)
    add_table(
        doc,
        "ตารางภาคผนวก ง.1 แหล่งอ้างอิงหลัก",
        ["Source", "URL / File", "Use"],
        [
            ["WHO", "who.int/news-room/fact-sheets/detail/cervical-cancer", "public-health framing"],
            ["CDC", "cdc.gov/cervical-cancer/screening", "Pap vs HPV test distinction"],
            ["FDA/IMDRF GMLP", "fda.gov/.../good-machine-learning-practice", "medical-AI lifecycle governance"],
            ["CONSORT-AI", "nature.com/articles/s41591-020-1034-x", "clinical AI reporting"],
            ["DECIDE-AI", "nature.com/articles/s41591-022-01772-9", "early clinical evaluation"],
            ["Project metrics", "models/*.json", "model-performance claims"],
        ],
        [3.5, 6.2, 6.0],
        "docs/SOURCE_CITATION_LEDGER.md",
    )

    add_heading(doc, "คำรับรองความซื่อสัตย์ของรายงาน", 1)
    add_bullets(doc, [
        "รายงานนี้ใช้ metrics จาก canonical JSON เท่านั้น",
        "ไม่อ้างว่า CerviCo-Pilot ตรวจพบ HPV infection",
        "ระบุชัดว่ายังขาด Thai-domain validation",
        "ระบุชัดว่ายังไม่ใช่ระบบสำหรับใช้งานคลินิกแบบอัตโนมัติ",
        "ระบุข้อจำกัด KOIL, SCC recall, calibration และ domain shift อย่างชัดเจน",
    ])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build_doc()
    print(out)
